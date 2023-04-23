
from docx import Document
import win32com.client as win32
from docx.shared import Inches
'''

def createWordFile():
    doc = Document()
    data = (
        ( "Фамилия", str(Surname.get())),
        ( "Пол",str(name_ManWom.get())),
        ("Дата анализа", str(name_dateAnal.get())),
        ("Возраст", str(name_age.get())),
        ( "Диагноз основной",str(name_DiagnosBasic.get())),
        ( "Диагноз сопутствующий",str(name_DiagnosRelated.get())),
        ("1.Гены", str(name_gens1.get())),
        ("2.Гены", str(name_2gens.get())),
        ("3.Гены", str(name_3gens.get())),
        ( "Сезон",str(name_Season.get())),

        ("РЕЗУЛЬАТЫ ГЕМОТОЛОГИЧЕСКОГО ИССЛЕДОВАНИЯ",""),
        ("Лейкоциты (WBC)", str(name_WBC.get())),
        ( "Лимфоциты (LYMF)",str(name_LYMF.get())),
        ("Моноциты (MON)", str(name_MON.get())),
        ("Нейтрофилы (NEU)", str(name_NEU.get())),
        ("Эозинофилы (EOS)", str(name_EOS.get())),
        ("Базофилы (BAS)", str(name_HGB.get())),
        ("Гемоглобин (HGB)", str(name_BAS.get())),
        ("Тромбоциты (PLT)", str(name_PLT.get())),

        ("ИМУННЫЙ СТАТУС", ""),
        ( "Общие T-лимфоциты (CD45+CD3+)",str(name_CD3.get())),
        ("Общие В-лимфоциты (CD45+CD19+)", str(name_CD19.get())),
        ( "Т-хелперы (CD45+CD3+CD4+)",str(name_CD4.get())),
        ("Соотношение CD3+CD4+/CD3+CD8+  ", str(name_ratio.get())),
        ( "Т-цитотоксические лимфоциты (CD45+CD3+СD8+)",str(name_CD8.get())),
        ("Циркулирующие иммунные комплексы", str(name_CIK.get())),
        ("Общие NK-клетки (CD45+CD3-CD16+56+) ", str(name_NK_Ob.get())),
        ("NK-клетки цитолитические (CD45+CD3-CD16brightCD56dim) ", str(name_NK_Cit.get())),
        ("HCI-тест(спонтанный)", str(name_HCI_CO.get())),
        ("HCI-тест(стимулированый)", str(name_HCI_CT.get())),

        ("ЦИТОКИНОВЫЙ СТАТУС", ""),
        ("CD3+IFNy+(стимулированный)", str(name_CD3IFN_CT.get())),
        ("CD3+IFNy+(спонтанный)", str(name_CD3INF_CO.get())),
       # ("Индекс (CD3+IFNy+(стимулированный)/CD3+IFNy+(спонтанный))", str(interferon.get())),

        ("CD3+TFNy+(стимулироанный)	 ", str(name_CD3INF_CO.get())),
        ("CD3+TFNy+(спонтанный)	 ", str(name_CD3INF_CO.get())),
        #("Индекс (CD3+TNFa+(стимулированный)/CD3+TNFa+(спонтанный))", str(fno.get())),

        ("CD3+IL2+(стимулированный)	 ", str(name_CD3IL2_CT.get())),
        ( "CD3+IL2+(спонтанный)	 ",str(name_CD3IL2_CO.get())),
        #("Индекс (CD3+IL2+(стимулированный)/CD3+IL2+(спонтанный))", str(interlikin.get())),
        ( "ФНО",str(name_FNO.get())),

    )
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    for id, name in data:
        row = table.add_row().cells
        row[0].text = str(id)
        row[1].text = name


    picture1 = doc.add_picture('file1.png', width=Inches(5))
    picture2 = doc.add_picture('file2.png', width=Inches(5))
    picture3 = doc.add_picture('file3.png', width=Inches(5))

    doc.save('test.docx')

'''
def SaveGraphicInPNG():
    excel = win32.gencache.EnsureDispatch('Excel.Application')
    workbook = excel.Workbooks.Open(r'C:\Users\artem\PycharmProjects\neural\Графики.xlsx')
    worksheet = workbook.Worksheets('grafik')
    c = 1
    for chart in worksheet.ChartObjects():
        file_name = f'file{c}.png'
        ch = worksheet.ChartObjects(chart.Name).Chart
        ch.Export(r'C:\Users\artem\PycharmProjects\neural' + f'\{file_name}', 'PNG')
        print(chart.Name)
        c += 1
    workbook.Close()
    excel.Quit()


SaveGraphicInPNG()