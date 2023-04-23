from docx import Document
import sqlite3 as sq
import win32com.client as win32
from docx.shared import Inches
from openpyxl.reader.excel import load_workbook
from openpyxl.workbook import Workbook
import win32api
import win32con
import win32gui
from PIL import Image, ImageEnhance
import getpass

conn = sq.connect('Пациенты1.db')
cursor = conn.cursor()
sqlite_select_query = """SELECT * from анализы"""
cursor.execute(sqlite_select_query)

surname = "4"
cursor.execute("SELECT * from анализы WHERE num="+f"{surname}")
result = cursor.fetchone()


def SaveGraphicInPNG(list_name):
    excel = win32.gencache.EnsureDispatch('Excel.Application')
    workbook = excel.Workbooks.Open(f'C:/Users/{getpass.getuser()}/PycharmProjects/neural/графики пациентов/' + f'{list_name}')
    worksheet = workbook.Worksheets('grafik')
    c = 1
    for chart in worksheet.ChartObjects():
        file_name = f'{list_name}' + ' ' + f'{c}.png'
        ch = worksheet.ChartObjects(chart.Name).Chart
        ch.Export('C:/Users/artem/PycharmProjects/neural/Графики/' + f'{file_name}', 'PNG')
        print(chart.Name)
        c += 1
    workbook.Close()
    excel.Quit()


def createWordFile(rows, pic_name):

    surname = rows[2]
    date = rows[4]
    doc = Document()
    arr = [rows[30], rows[31], rows[32], rows[33], rows[34], rows[35]]
    res = []
    for i in arr:
        if i is None:
            res.append(1)
        else:
            res.append(i)
    print(res)
    interferon = float(res[0])/float(res[1])
    fno = float(res[2]) / float(res[3])
    interlikin = float(res[4]) / float(res[5])

    data = (
        ("Фамилия", str(rows[2])),
        ("Пол", str(rows[3])),
        ("Дата анализа", str(rows[4])),
        ("Возраст", str(rows[5])),
        ("Диагноз основной", str(rows[6])),
        ("Диагноз сопутствующий", str(rows[7])),
        ("1.Гены", str(rows[8])),
        ("2.Гены", str(rows[9])),
        ("3.Гены", str(rows[10])),
        ("Сезон", str(rows[11])),

        ("РЕЗУЛЬАТЫ ГЕМОТОЛОГИЧЕСКОГО ИССЛЕДОВАНИЯ", ""),
        ("Лейкоциты (WBC)", str(rows[12])),
        ("Лимфоциты (LYMF)", str(rows[13])),
        ("Моноциты (MON)", str(rows[14])),
        ("Нейтрофилы (NEU)", str(rows[15])),
        ("Эозинофилы (EOS)", str(rows[18])),
        ("Базофилы (BAS)", str(rows[19])),
        ("Гемоглобин (HGB)", str(rows[16])),
        ("Тромбоциты (PLT)", str(rows[17])),

        ("ИМУННЫЙ СТАТУС", ""),
        ("Общие T-лимфоциты (CD45+CD3+)", str(rows[20])),
        ("Общие В-лимфоциты (CD45+CD19+)", str(rows[21])),
        ("Т-хелперы (CD45+CD3+CD4+)", str(rows[22])),
        ("Соотношение CD3+CD4+/CD3+CD8+  ", str(rows[24])),
        ("Т-цитотоксические лимфоциты (CD45+CD3+СD8+)", str(rows[23])),
        ("Циркулирующие иммунные комплексы", str(rows[27])),
        ("Общие NK-клетки (CD45+CD3-CD16+56+) ", str(rows[26])),
        ("NK-клетки цитолитические (CD45+CD3-CD16brightCD56dim) ", str(rows[25])),
        ("HCI-тест(спонтанный)", str(rows[28])),
        ("HCI-тест(стимулированый)", str(rows[29])),

        ("ЦИТОКИНОВЫЙ СТАТУС", ""),
        ("CD3+IFNy+(стимулированный)", str(rows[30])),
        ("CD3+IFNy+(спонтанный)", str(rows[31])),
        ("Индекс (CD3+IFNy+(стимулированный)/CD3+IFNy+(спонтанный))", str(interferon)),

        ("CD3+TFNy+(стимулироанный)	 ", str(rows[32])),
        ("CD3+TFNy+(спонтанный)	 ", str(rows[33])),
        ("Индекс (CD3+TNFa+(стимулированный)/CD3+TNFa+(спонтанный))", str(fno)),

        ("CD3+IL2+(стимулированный)	 ", str(rows[34])),
        ("CD3+IL2+(спонтанный)	 ", str(rows[35])),
        ("Индекс (CD3+IL2+(стимулированный)/CD3+IL2+(спонтанный))", str(interlikin)),
        ("ФНО", rows[36]),
    )
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    for id, name in data:
        row = table.add_row().cells
        row[0].text = str(id)
        row[1].text = name

    doc.add_paragraph('График T-клеточного звена')
    picture1 = doc.add_picture(f'C:/Users/{getpass.getuser()}/PycharmProjects/neural/Графики/{pic_name} 1.png', width=Inches(5))
    doc.add_paragraph('График B-клеточного звена')
    picture2 = doc.add_picture(f'C:/Users/{getpass.getuser()}/PycharmProjects/neural/Графики/{pic_name} 2.png', width=Inches(5))
    doc.add_paragraph('График цитокиновых пар')
    picture3 = doc.add_picture(f'C:/Users/{getpass.getuser()}/PycharmProjects/neural/Графики/{pic_name} 3.png', width=Inches(5))
    print("ааа")

    if date == 'None':
        date = ''

    doc.save(f'C:/Users/{getpass.getuser()}/PycharmProjects/neural/Печатные Формы/' + f'{surname}{date}.docx')


def check(max, min, value):
    if (value > max):
        return "- Отклонение от нормы больше 20% нормы вверх"
    if (value < min):
        return "- Отклонение от нормы больше 20% нормы вниз"
    if (value < min and value > max):
        return "- В пределах нормы"

def createRecomendation(rows):
    arr = [rows[30], rows[31], rows[32], rows[33], rows[34], rows[35]]
    res = []
    for i in arr:
        if i is None:
            res.append(1)
        else:
            res.append(i)
    interferon = round(float(res[0]) / float(res[1]), 2)
    fno = round(float(res[2]) / float(res[3]), 2)
    interlikin = round(float(res[4]) / float(res[5]), 2)

    neu_lymf = round (float(rows[15]) / float(rows[13]) ,2)
    neu_cd3 = round (float(rows[15]) / float(rows[20]) ,2)
    neu_cd4 = round (float(rows[15]) / float(rows[22]) ,2)
    neu_cd8 = round (float(rows[15]) / float(rows[23]) ,2)
    lymf_cd19 = round (float(rows[13]) / float(rows[21]) ,2)
    cd19_cd4 = round (float(rows[21]) / float(rows[22]) ,2)
    cd19_cd8 = round (float(rows[21])/ float(rows[23]) ,2)

    normal_values_max = [1.80, #NEU/LYMF
                         3.70, #NEU/CD3
                         12.30,#NEU/CD4
                         5.00, #NEU/CD8
                         10.00,#LYMF/CD19
                         0.80, #CD19/CD4
                         0.30,#CD19/CD8
                         120, #FNO
                         120, #INTERLIKIN
                         120] #INTERFERON

    normal_values_min = [1.67, #NEU/LYMF
                         2.30, #NEU/CD3
                         9.50, #NEU/CD4
                         3.00, #NEU/CD8
                         9.60, #LYMF/CD19
                         0.50, #CD19/CD4
                         0.20, #CD19/CD8
                         80,  # FNO
                         80,  # INTERLIKIN
                         80]  # INTERFERON

    document = Document()

    document.add_paragraph('Показатели T-клеточного иммунитета')
    document.add_paragraph('     NEU/LYMF ' f'{check(normal_values_max[0],normal_values_min[0],neu_lymf)}')
    document.add_paragraph('     NEU/CD3 'f'{check(normal_values_max[1],normal_values_min[1],neu_cd3)}')
    document.add_paragraph('     NEU/CD4 'f'{check(normal_values_max[2],normal_values_min[2],neu_cd4)}')
    document.add_paragraph('     NEU/CD8 'f'{check(normal_values_max[3],normal_values_min[3],neu_cd8)}')
    print("значения", "минимум", "максимум")
    print(neu_lymf, normal_values_min[0], normal_values_max[0])
    print(neu_cd3, normal_values_min[1], normal_values_max[1])
    print(neu_cd4, normal_values_min[2], normal_values_max[2])
    print(neu_cd8, normal_values_min[3], normal_values_max[3])

    document.add_paragraph('Показатели B-клеточного иммунитета')
    document.add_paragraph('     LYMF/CD19 'f'{check(normal_values_max[4],normal_values_min[4],lymf_cd19)}')
    document.add_paragraph('     CD19/CD4 'f'{check(normal_values_max[5],normal_values_min[5],cd19_cd4)}')
    document.add_paragraph('     CD19/CD8 'f'{check(normal_values_max[6],normal_values_min[6],cd19_cd8)}')

    print(lymf_cd19, normal_values_min[4], normal_values_max[4])
    print(cd19_cd4, normal_values_min[5], normal_values_max[5])
    print(cd19_cd8, normal_values_min[6], normal_values_max[6])

    document.add_paragraph('Цитокиновые пары')
    document.add_paragraph('     ФНО 'f'{check(normal_values_max[7],normal_values_min[7],fno)}')
    document.add_paragraph('     ИНТЕРФЕРОН 'f'{check(normal_values_max[8],normal_values_min[8],interferon)}')
    document.add_paragraph('     ИНТЕРЛЕКИН 'f'{check(normal_values_max[9],normal_values_min[9],interlikin)}')

    print(fno, normal_values_min[7], normal_values_max[7])
    print(interferon, normal_values_min[8], normal_values_max[8])
    print(interlikin, normal_values_min[9], normal_values_max[9])

    document.save('example.docx')


def Graphic(rows, list):

    interferon = float(rows[30]) / float(rows[31])
    fno = float(rows[33]) / float(rows[32])
    interlikin = float(rows[34]) / float(rows[35])

    # создание графиков
    wb = Workbook()
    ws = wb.active
    curr_wb = load_workbook(f'C:/Users/{getpass.getuser()}/PycharmProjects/neural/графики пациентов/' + f'{list}')
    ws = curr_wb['grafik']

    ws['I22'] = "Значения"
    ws['I23'] = "ФНО"
    ws['I24'] = "Инерликин"
    ws['I25'] = "Интерферон"

    ws['J22'] = "Максимальные значения"
    ws['K22'] = "Минимальные значения"
    ws['L22'] = "Значения пациента"
    ws['M22'] = "Названия"

    ws['J23'] = 120
    ws['J24'] = 120
    ws['J25'] = 120

    ws['K23'] = 60
    ws['K24'] = 60
    ws['K25'] = 60

    ws['L23'] = fno
    ws['L24'] = interlikin
    ws['L25'] = interferon

    curr_wb.save(f'C:/Users/{getpass.getuser()}/PycharmProjects/neural/графики пациентов/' + f'{list}')
    curr_wb.close()


conn = sq.connect('Пациенты1.db')
cursor = conn.cursor()
sqlite_select_query = """SELECT * from анализы"""
cursor.execute(sqlite_select_query)
records = cursor.fetchall()


def get_data(row):
    data = f'{row[2]} {row[4]}.xlsx'
    return data


for i in records:
    print(get_data(i))
    SaveGraphicInPNG(get_data(i))
    createWordFile(i, get_data(i))
    createRecomendation(i)
#Graphic(result, get_data(i))

cursor.close()
conn.close()
